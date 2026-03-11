"""
Document loaders — one function per format, all returning raw text + metadata.
"""

import json
import re
import pathlib
from typing import Tuple

def load_docx(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required to load .docx files. Install with: pip install python-docx")

    doc = Document(path)
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append("  ".join(cells))
    return "\n".join(lines)

def load_markdown(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()

def load_txt(path: str) -> str:
    with open(path, encoding="utf-8", errors="replace") as f:
        return f.read()

def load_json(path: str) -> str:
    """Flatten JSON to text by recursively extracting string values."""
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    return _flatten_json(data)


def _flatten_json(obj, prefix="", lines=None) -> str:
    if lines is None:
        lines = []
    if isinstance(obj, dict):
        for k, v in obj.items():
            full_key = f"{prefix}.{k}" if prefix else k
            if isinstance(v, str):
                lines.append(f"{full_key}: {v}")
            else:
                _flatten_json(v, full_key, lines)
    elif isinstance(obj, list):
        for i, item in enumerate(obj):
            _flatten_json(item, f"{prefix}[{i}]", lines)
    elif isinstance(obj, (int, float, bool)):
        lines.append(f"{prefix}: {obj}")
    elif obj is not None:
        lines.append(f"{prefix}: {obj}")
    return "\n".join(lines)

def load_yaml(path: str) -> str:
    try:
        import yaml
    except ImportError:
        raise ImportError("PyYAML is required to load .yaml/.yml files. Install with: pip install pyyaml")
    with open(path, encoding="utf-8") as f:
        data = yaml.safe_load(f)
    return _flatten_json(data)

def load_csv(path: str) -> str:
    import csv
    lines = []
    with open(path, encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f)
        for row in reader:
            cleaned = [cell.strip() for cell in row if cell.strip()]
            if cleaned:
                lines.append("  ".join(cleaned))
    return "\n".join(lines)

def load_pdf(path: str) -> str:
    """Extract text from text-based PDFs. Requires pdfminer or pypdf."""
    try:
        import pypdf
        reader = pypdf.PdfReader(path)
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n".join(pages)
    except ImportError:
        pass

    try:
        from pdfminer.high_level import extract_text
        return extract_text(path)
    except ImportError:
        pass

    raise ImportError(
        "A PDF parsing library is required. Install one with:\n"
        "  pip install pypdf\n"
        "  -- or --\n"
        "  pip install pdfminer.six"
    )

LOADERS = {
    ".docx": load_docx,
    ".doc":  load_docx,
    ".md":   load_markdown,
    ".txt":  load_txt,
    ".json": load_json,
    ".yaml": load_yaml,
    ".yml":  load_yaml,
    ".csv":  load_csv,
    ".pdf":  load_pdf,
}


def load_document(path: str) -> str:
    """Auto-detect format and return extracted text."""
    ext = pathlib.Path(path).suffix.lower()
    loader = LOADERS.get(ext)
    if loader is None:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {', '.join(sorted(LOADERS))}"
        )
    return loader(path)


def supported_extensions() -> list:
    return sorted(LOADERS.keys())
